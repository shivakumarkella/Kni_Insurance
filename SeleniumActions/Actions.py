# Wait for Element and Return Element
# Screenshots
# Write a method Sleep()
# Writ a method to select * Dropdown, * check box, * radio Button * Alerts * Warnings


from pytz import timezone
from tzlocal import get_localzone
import os as OS
import docx as d
from selenium import webdriver
from selenium.webdriver.common.by import By as BY
from selenium.webdriver.support.ui import Select
from datetime import datetime
import time
from selenium.webdriver.support.ui import WebDriverWait as wdw
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import  *
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys
from TestResults.customLogger import customLogger
import logging
format = "%Y_%m_%d_%H_%M_%S" # time stamp will come in this format


class SeleniumActions():
    def __init__(self,driver):
        self.driver = driver
        self.timeStamp = (datetime.now(timezone(str(get_localzone()).upper()))).strftime(format)

    def by_type(self,locatortype):
        if locatortype.lower()=="id":
            return BY.ID
        elif locatortype.lower()=="xpath":
            return BY.XPATH
        elif locatortype.lower() =="name":
            return BY.NAME
        elif locatortype.lower() =="link_text":
            return BY.LINK_TEXT
        elif locatortype.lower() == "class_name":
            return BY.CLASS_NAME
        elif locatortype.lower() == "tag_name":
            return BY.TAG_NAME
        elif locatortype.lower() == "css_selector":
            return BY.CSS_SELECTOR
        elif locatortype.lower() == "partial_link_text":
            return BY.PARTIAL_LINK_TEXT
        else:
            print("Locator Type is not supported " + locatortype)

    def get_element(self, locator, locatortype="id"):
         byType = self.by_type(locatortype)
         element = self.driver.find_element(byType, locator)
         return element

    def sendKeys(self, message,locator=None, locatortype=None, element=None):
        if element is None :
            element = self.get_element(locator=locator, locatortype=locatortype)
        element.send_keys(message)


    def timetosleep(self, timetosleep=10):
        time.sleep(timetosleep)


    def waitForActionsOnElement(self,timeout=10, pollFrequency=0.5):
        try:
            wait = wdw(self.driver, timeout=timeout, poll_frequency=pollFrequency,
                                 ignored_exceptions=[NoSuchElementException,
                                                     ElementNotVisibleException,
                                                     ElementNotSelectableException,TimeoutException])

            # element = wait.until(EC.element_to_be_clickable((byType, locator)))

        except:
            print('Error while wait for the element.')
        return wait

    def waitForElementToDisappear(self,locator,locatorType='id',timeout=10,element=None):
            if element is None:
                element=self.get_element(locator=locator,locatortype=locatorType)
            wait=self.waitForActionsOnElement(timeout=timeout,pollFrequency=1)
            wait.until(EC.invisibility_of_element_located(element))


    def waitForElementToVisible(self,locator,locatorType='id',timeout=10,element=None):
        try:
            if element is None:
                element=self.get_element(locator=locator,locatortype=locatorType)
            wait=self.waitForActionsOnElement(timeout=timeout,pollFrequency=1)
            wait.until(EC.visibility_of_element_located(element))
        except:
            print('Error while waiting for the element')
        return element





    def waitForElementToClickOn(self,locator,locatorType='id',timeout=10):
        try:
            element = None
            wait=self.waitForActionsOnElement(timeout=timeout,pollFrequency=1)
            byType=self.by_type(locatortype=locatorType)
            element=wait.until(EC.element_to_be_clickable((byType,locator)))
            if element is not None:
               element.click()

        except:
            print('Waited for' + locator + 'to be clicked, but ' + locatorType + 'not found')


    def isElementPresent(self,locator,locatorType='id',timeout=10):
        wait=self.waitForActionsOnElement()
        byType=self.by_type(locatortype=locatorType)
        element=wait.until(EC.presence_of_element_located((byType,locator)))
        if element is not None:
            return True
        else:
            return False




    def selectDropdown(self,locator, locatorType,element=None,SelectValue=None):
        element = Select(self.get_element(locatortype=locatorType,
                                           locator=locator))
        self.log = customLogger(logging.DEBUG)
        self.log.error(SelectValue)
        element.select_by_value(SelectValue)


    



    def selectRadiobutton(self,locator,locatorType,element=None):
        if element is None:
            element = self.get_element(locator=locator, locatortype=locatorType)
            element.click()
            element.click()
        else:
            element.click()


    def selectCheckbox(self,locator,locatorType,element=None):
        if element is None:
            element =self.get_element(locator=locator, locatortype=locatorType)
            element.click()
        else:
            element.click()


    def handlingalert(self):
       checkalert= self.driver.switch_to.alert
       checkalert.accept()


    def waitForPageLoaded(self):
        wait = self.waitForActionsOnElement(timeout=30)
        return str((self.driver).execute_script("return document.readyState")) == "complete"

    def LoadtheDomCompletely(self):
        try:
            i = 1
            while (i == 1):
                isLoaded = self.waitForPageLoaded()
                if isLoaded == True:
                    i = 0
        except:
            print('time out , page not loaded yet')

    def tabout(self):
       A= ActionChains(driver=self.driver)
       A.send_keys(Keys.TAB)


    def getTheDropDownSelectedValue(self,locator, locatorType,element=None,SelectValue=None):
        selectedValue = Select(self.get_element(locatortype=locatorType,
                                          locator=locator)).first_selected_option.get_attribute("value")
        return selectedValue

    def getListOFWindows(self):
        handles=self.driver.window_handles
        return handles

    def sleepForWhile(self,sleepTime=5):
        time.sleep(sleepTime)


    def clickmethod(self,locator, locatorType):
        element = self.get_element(locator=locator,locatortype=locatorType)
        element.click()

    def getTheCurrentWindowSize(self):
        # get the Window size i.e. height and width
            height = self.driver.execute_script("return window.innerHeight;")
            width = self.driver.execute_script("return window.innerWidth;")
            return height,width


    def scrollDown(self,height=0,width=0):
        if height==0 or width==0:
            height,width=self.getTheCurrentWindowSize()
        scriptForScroll = 'window.scrollBy(0,' + str(height) + ');'
        self.driver.execute_script("{}".format(scriptForScroll))

    def scrollUp(self,height=0,width=0):
        if height==0 or width==0:
            height,width=self.getTheCurrentWindowSize()
        scriptForScroll = 'window.scrollBy(0,' + str(-height) + ');'
        self.driver.execute_script("{}".format(scriptForScroll))

    def highlight(self, element,directory_to_save_ScreenShot='Default'):
        """Highlights (blinks) a Selenium Webdriver element"""
        driver = self.driver
        def apply_style(s):
            driver.execute_script("arguments[0].setAttribute('style', arguments[1]);",element, s)
        original_style = element.get_attribute('style')
        apply_style("background: yellow; border: 2px solid red;")
        if directory_to_save_ScreenShot.lower()=='default':
            print('only highlight the web element')
        else:
            self.driver.save_screenshot(directory_to_save_ScreenShot)
        time.sleep(1)
        apply_style(original_style)




    def takeScreenShot(self, element='Default', typeOfScreenShot='Default', screenShotName='Default', path='Default',
                       returnScreenShotName=False):
        typeOfScreenShot = typeOfScreenShot.lower()
        screenShotName = screenShotName.lower()
        path = path.lower()
        if typeOfScreenShot == 'default':
            typeOfScreenShot = '.png'
        else:
            typeOfScreenShot = '.jpg'
        if screenShotName == 'default':
            screenShotName = 'Demo'
        if path == 'default':
            path = 'C:\\Users\\kella\\workspace_python\\Kni_Insurance\\TestResults\\DocsScreenShotsLogFiles'

        # To maintain Unique file name adding the time stamp to the screenshot
        screenShotName = str(screenShotName) + str(self.timeStamp)
        # by adding timestamp (.) is part of the file name, hence replace (.) with empty
        screenShotName = screenShotName.replace('.', '')

        directory_to_save_ScreenShot = path + screenShotName + typeOfScreenShot

        if str(element).lower() == 'default':
            self.driver.save_screenshot(directory_to_save_ScreenShot)
        else:
            self.highlight(element, directory_to_save_ScreenShot)

        if returnScreenShotName == True:
            return directory_to_save_ScreenShot


    def saveScreenShotInDocxFile(self,locator='Default',locatorType='Default',fileName='Default',
                                 heading='Default',width=0,height=0,deleteScreenShot='Default',
                                 paragraphTobeAdded='Default',path='Default',addScreenShotToSameDoc='Default',
                                 returnDocName=False):
        if fileName.lower()=='default':
            fileName='Demo'
        if width==0:
            #Set Default Width as 6 Inches
            width=6
        if height==0:
            #set Default Height as 6 inches
            height=6
        if path.lower()=='default':
            path='C:\\Users\\kella\\workspace_python\\Kni_Insurance\\TestResults\\DocsScreenShotsLogFiles\\'
        if addScreenShotToSameDoc.lower()=='default':
            # by default we will save the screen shots in the new file
            addScreenShotToSameDoc=False
            # To maintain Unique file name adding the time stamp to the Document
            fileName = str(fileName) + str(self.timeStamp)
            # add the extension to the file , word document .docx
            fileName = fileName + '.docx'
        # #neew code for the Element access

        # element=Object_element.waitForElement(locatorType,locator)
        if locatorType.lower()=='default' or locator.lower()=='default':
            screenShotPath = self.takeScreenShot(returnScreenShotName=True)
        else:

            element=self.get_element(locator=str(locator),locatortype=str(locatorType))
            screenShotPath = self.takeScreenShot(element,returnScreenShotName=True)

        directoryTosavetheDocFile=path+fileName
        #check the document exist  in the directory or not
        is_document_exist=OS.path.exists(directoryTosavetheDocFile)
        if is_document_exist==True:
            #using the exisitng document
            document=d.Document(directoryTosavetheDocFile)
            document.add_page_break()
        else:
            document=d.Document()
        if heading.lower() == 'default':
            heading = ''
        else:
            document.add_heading(heading,0)

        if paragraphTobeAdded.lower()=='default':
            #nothing is there to add into the document
            paragraphTobeadded=False
        else:
            document.add_paragraph(paragraphTobeAdded)

        document.add_picture(screenShotPath,width=d.shared.Inches(width),height=d.shared.Inches(height))

        if deleteScreenShot.lower()=='default':
            #by default we will delete the screen shot once we paste the it into the Word documnet
            deleteScreenShot=True
            OS.remove(screenShotPath)
        document.save(directoryTosavetheDocFile)

        if returnDocName==True:
            return fileName

